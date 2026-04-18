from __future__ import annotations

import re
import secrets
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from bson import ObjectId
from docx import Document

from app.config import settings
from app.services.oss import OssPutResult, put_file_from_local


EXPORTS_TMP_DIR = Path("/tmp/uploads/exports")
EXPORTS_TMP_DIR.mkdir(parents=True, exist_ok=True)


def _now() -> datetime:
    return datetime.utcnow()


def _dt_to_iso(v: datetime) -> str:
    return v.replace(tzinfo=timezone.utc).isoformat().replace("+00:00", "Z")


def _safe_filename_part(s: str) -> str:
    s = (s or "").strip() or "未命名"
    # remove path separators and control chars
    s = re.sub(r"[\\/\x00-\x1f]", "_", s)
    # limit length
    return s[:60]


def build_export_filename(*, title: str, now: Optional[datetime] = None) -> str:
    now = now or _now()
    date_s = now.strftime("%Y-%m-%d")
    title_part = _safe_filename_part(title)
    return f"成长档案_{date_s}_{title_part}.docx"


def render_story_to_docx_bytes(*, story: dict[str, Any]) -> bytes:
    """Render story JSON into a DOCX file.

    Minimal, stable formatting for MVP.
    """

    title = str(story.get("title") or "成长档案").strip() or "成长档案"
    doc = Document()
    doc.add_heading(title, level=0)

    profile = story.get("profile") if isinstance(story, dict) else None
    if isinstance(profile, dict):
        name = str(profile.get("name") or "").strip()
        goal = str(profile.get("learningGoal") or "").strip()
        pref = str(profile.get("preferences") or "").strip()
        if name or goal or pref:
            doc.add_heading("档案资料", level=1)
            if name:
                doc.add_paragraph(f"姓名：{name}")
            if goal:
                doc.add_paragraph(f"学习目标：{goal}")
            if pref:
                doc.add_paragraph(f"学习偏好：{pref}")

    chapters = story.get("chapters") or []
    if isinstance(chapters, list):
        for idx, ch in enumerate(chapters, start=1):
            if not isinstance(ch, dict):
                continue
            ch_title = str(ch.get("chapterTitle") or f"章节{idx}")
            doc.add_heading(ch_title, level=1)
            paras = ch.get("paragraphs") or []
            if isinstance(paras, list):
                for p in paras:
                    if p is None:
                        continue
                    doc.add_paragraph(str(p))

            # optional timeline
            tl = ch.get("timeline") or []
            if isinstance(tl, list) and tl:
                doc.add_paragraph("时间线：")
                for it in tl:
                    if not isinstance(it, dict):
                        continue
                    at = str(it.get("at") or "")
                    summary = str(it.get("summary") or "")
                    doc.add_paragraph(f"- {at} {summary}")

    milestones = story.get("milestones") or []
    if isinstance(milestones, list) and milestones:
        doc.add_heading("里程碑", level=1)
        for ms in milestones:
            if not isinstance(ms, dict):
                continue
            at = str(ms.get("at") or "")
            m_title = str(ms.get("title") or "")
            summary = str(ms.get("summary") or "")
            doc.add_paragraph(f"- {at} {m_title}：{summary}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def render_story_to_pdf_bytes(*, story: dict[str, Any]) -> bytes:
    """Render story JSON into a simple PDF.

    This is not a docx->pdf conversion. It's a direct render for preview.
    """

    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    # Register a CJK-capable font (built-in in reportlab) for Chinese.
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
    except Exception:
        font_name = "Helvetica"

    title = str(story.get("title") or "成长档案").strip() or "成长档案"

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    margin_x = 48
    y = h - 64
    line_h = 18

    def _new_page():
        nonlocal y
        c.showPage()
        c.setFont(font_name, 12)
        y = h - 64

    def _draw_line(text: str, *, size: int = 12):
        nonlocal y
        if y < 64:
            _new_page()
        c.setFont(font_name, size)
        # naive wrap by char count (good enough for MVP)
        t = (text or "").strip()
        if not t:
            y -= line_h
            return
        max_chars = 45
        for i in range(0, len(t), max_chars):
            if y < 64:
                _new_page()
            c.drawString(margin_x, y, t[i : i + max_chars])
            y -= line_h

    # title
    _draw_line(title, size=18)
    y -= 6

    profile = story.get("profile") if isinstance(story, dict) else None
    if isinstance(profile, dict):
        name = str(profile.get("name") or "").strip()
        goal = str(profile.get("learningGoal") or "").strip()
        pref = str(profile.get("preferences") or "").strip()
        if name or goal or pref:
            _draw_line("档案资料", size=14)
            if name:
                _draw_line(f"姓名：{name}", size=12)
            if goal:
                _draw_line(f"学习目标：{goal}", size=12)
            if pref:
                _draw_line(f"学习偏好：{pref}", size=12)
            y -= 6

    chapters = story.get("chapters") or []
    if isinstance(chapters, list):
        for idx, ch in enumerate(chapters, start=1):
            if not isinstance(ch, dict):
                continue
            ch_title = str(ch.get("chapterTitle") or f"章节{idx}")
            _draw_line(ch_title, size=14)
            paras = ch.get("paragraphs") or []
            if isinstance(paras, list):
                for p in paras:
                    _draw_line(str(p), size=12)
            y -= 6

    milestones = story.get("milestones") or []
    if isinstance(milestones, list) and milestones:
        _draw_line("里程碑", size=14)
        for ms in milestones:
            if not isinstance(ms, dict):
                continue
            at = str(ms.get("at") or "")
            m_title = str(ms.get("title") or "")
            summary = str(ms.get("summary") or "")
            _draw_line(f"- {at} {m_title}：{summary}")

    c.save()
    return buf.getvalue()


@dataclass(frozen=True)
class ExportCreateResult:
    export_id: str
    file_name: str
    object_key: str
    pdf_file_name: str
    pdf_object_key: str
    created_at: datetime


def create_export_to_oss(*, story_id: str, user_id: str, account_id: str, story_doc: dict[str, Any], template: str = "default") -> ExportCreateResult:
    """Create an export snapshot for a story and upload to OSS.

    Returns metadata; caller is responsible for storing it in DB.
    """

    created_at = _now()
    export_id = f"e_{ObjectId()}"

    story = story_doc.get("story") if isinstance(story_doc, dict) else None
    if not isinstance(story, dict):
        raise ValueError("story doc missing story")

    title = str(story.get("title") or "成长档案")
    file_name = build_export_filename(title=title, now=created_at)
    pdf_file_name = file_name[:-5] + ".pdf" if file_name.lower().endswith(".docx") else file_name + ".pdf"

    content = render_story_to_docx_bytes(story=story)

    pdf_content = render_story_to_pdf_bytes(story=story)

    # Write to a tmp file (put_file_from_local expects a Path)
    tmp_fp = EXPORTS_TMP_DIR / f"{export_id}.docx"
    tmp_fp.write_bytes(content)
    tmp_pdf_fp = EXPORTS_TMP_DIR / f"{export_id}.pdf"
    tmp_pdf_fp.write_bytes(pdf_content)

    # Object key layout
    object_key = f"exports/{user_id}/{account_id}/{export_id}.docx"
    pdf_object_key = f"exports/{user_id}/{account_id}/{export_id}.pdf"

    put_res: OssPutResult = put_file_from_local(tmp_fp, object_key=object_key)
    put_pdf_res: OssPutResult = put_file_from_local(tmp_pdf_fp, object_key=pdf_object_key)

    # best-effort cleanup
    try:
        tmp_fp.unlink(missing_ok=True)  # type: ignore[call-arg]
    except Exception:
        pass
    try:
        tmp_pdf_fp.unlink(missing_ok=True)  # type: ignore[call-arg]
    except Exception:
        pass

    return ExportCreateResult(
        export_id=export_id,
        file_name=file_name,
        object_key=put_res.object_key,
        pdf_file_name=pdf_file_name,
        pdf_object_key=put_pdf_res.object_key,
        created_at=created_at,
    )


def generate_share_id() -> str:
    # 24 urlsafe chars ~ 144 bits entropy; with prefix it remains non-enumerable.
    return "Sh_" + secrets.token_urlsafe(24)
